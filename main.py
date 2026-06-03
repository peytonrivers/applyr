from fastapi import FastAPI, Depends, HTTPException, Request, Form, File, UploadFile, Depends
from fastapi.middleware.cors import CORSMiddleware
from jose import JWTError, jwt
from auth import router as auth_router
from dotenv import load_dotenv
import os
from database.database import Users, get_db
from database.storage import supabase
from pydantic import BaseModel
from sqlalchemy.orm import Session
import traceback



import json
import fitz  # PyMuPDF
from docx import Document #Docx
import io

load_dotenv()

SECRET_KEY = os.getenv("SECRET_KEY")
ALGORITHM = "HS256"

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://apply-r.com",
        "https://www.apply-r.com",
        "https://www.api.apply-r.com",
        "https://api.apply-r.com",
        "https://applyr-one.vercel.app",
        "https://applyr-frontend.onrender.com",
        "https://applyr-12k0.onrender.com",
        "http://localhost:5500",
        "http://127.0.0.1:5500",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.include_router(auth_router)


def get_current_user_id(request: Request) -> str:
    token = request.cookies.get("access_token")
    if not token:
        raise HTTPException(status_code=401, detail="No access token")
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = payload.get("sub")
        if user_id is None:
            raise HTTPException(status_code=401, detail="Invalid token")
        return user_id
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid token")


@app.get("/")
def root():
    return {"message": "ApplyR API is live"}


class SignupForm(BaseModel):
    first_name: str
    last_name: str
    phone_number: str
    us_citizen: str


@app.post("/complete-signup")
def complete_signup(
    data: SignupForm,
    user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db),
):
    user = db.query(Users).filter(Users.user_id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    try:
        print(data)
        user.first_name = data.first_name
        user.last_name = data.last_name
        user.phone_number = data.phone_number
        user.us_citizen = data.us_citizen
        db.commit()
        db.refresh(user)

        return {
            "message": "Signup completed successfully",
            "user_id": user.user_id,
            "email": user.email,
        }

    except Exception as e:
        print("COMPLETE SIGNUP ERROR:", repr(e))
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")
    




def extract_pdf_text(file_bytes: bytes) -> str:
    pdf = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""

    for page in pdf:
        text += page.get_text()

    pdf.close()
    return text.strip()

def extract_docx_text(file_bytes):
    doc = Document(io.BytesIO(file_bytes))

    text = []

    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    return "\n".join(text)



def extract_pdf_text(file_bytes: bytes) -> str:
    pdf = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""

    for page in pdf:
        text += page.get_text()

    pdf.close()
    return text.strip()


def extract_docx_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    text = []

    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    return "\n".join(text).strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    filename = filename.lower()

    if filename.endswith(".pdf"):
        return extract_pdf_text(file_bytes)

    if filename.endswith(".docx"):
        return extract_docx_text(file_bytes)

    raise HTTPException(
        status_code=400,
        detail="Only PDF and DOCX files are supported"
    )


@app.post("/complete-skill")
async def complete_skill(
    work_experience: str | None = Form(None),
    school: str | None = Form(None),
    resume: UploadFile = File(...),
    cover_letter: UploadFile = File(...),
    user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    user = db.query(Users).filter(Users.user_id == user_id).first()
    print(work_experience)
    print(school)
    print(resume)
    print(cover_letter)

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    try:
        work_experience_data = json.loads(work_experience) if work_experience else []
        school_data = json.loads(school) if school else []

        resume_bytes = await resume.read()
        cover_letter_bytes = await cover_letter.read()

        resume_text = extract_text(resume_bytes, resume.filename)
        cover_letter_text = extract_text(cover_letter_bytes, cover_letter.filename)

        resume_path = f"{user_id}/resume/{resume.filename}"
        cover_letter_path = f"{user_id}/cover-letter/{cover_letter.filename}"

        supabase.storage.from_("user-files").upload(
            path=resume_path,
            file=resume_bytes,
        )

        supabase.storage.from_("user-files").upload(
            path=cover_letter_path,
            file=cover_letter_bytes,
        )

        user.work_experience = work_experience_data
        user.school = school_data

        user.resume = resume_path
        user.resume_text = resume_text

        user.cover_letter = cover_letter_path
        user.cover_letter_text = cover_letter_text

        db.commit()
        db.refresh(user)

        return {
            "message": "Skill information saved successfully",
            "resume_path": resume_path,
            "cover_letter_path": cover_letter_path
        }

    except Exception as e:
        traceback.print_exc()
        db.rollback()
        raise HTTPException(
            status_code=500,
            detail=f"Database error: {str(e)}"
        )

class DetailsForm(BaseModel):
    country: str
    state: str
    city: str
    zipcode: str
    us_authorization: str
    disability: str
    veteran: str
    gender: str
    ethnicity: list[str]


@app.post("/complete-details")
def process_details(
    data: DetailsForm,
    user_id: str = Depends(get_current_user_id),
    db: Session = Depends(get_db)
):
    user = db.query(Users).filter(Users.user_id == user_id).first()

    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    try:
        user.country = data.country
        user.user_state = data.state
        user.city = data.city
        user.zipcode = data.zipcode
        user.us_work_authorization = data.us_authorization
        user.disability = data.disability
        user.veteran = data.veteran
        user.gender = data.gender
        user.ethnicity = data.ethnicity
        user.signup_complete = True

        db.commit()
        db.refresh(user)

        return {
            "message": "Details saved successfully",
            "user_id": user.user_id
        }

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))